from pathlib import Path
import json,re
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from PIL import Image
R=Path(__file__).resolve().parents[2]; data=json.loads((R/'artifacts/report-build/content.json').read_text())
doc=Document();sec=doc.sections[0];sec.page_width=Inches(8.27);sec.page_height=Inches(11.69);sec.top_margin=Inches(.72);sec.bottom_margin=Inches(.68);sec.left_margin=sec.right_margin=Inches(.74);sec.header_distance=Inches(.3);sec.footer_distance=Inches(.3)
for name in ['Normal','Title','Subtitle','Heading 1','Heading 2','Heading 3','Caption','Header','Footer']:
 st=doc.styles[name];st.font.name='Arial';st.font.color.rgb=RGBColor.from_string('000000');st.paragraph_format.space_after=Pt(9)
 st.font.size=Pt(11)
doc.styles['Normal'].paragraph_format.line_spacing=1.16
doc.styles['Title'].font.size=Pt(32);doc.styles['Title'].font.bold=True
doc.styles['Subtitle'].font.size=Pt(16)
doc.styles['Heading 1'].font.size=Pt(23);doc.styles['Heading 1'].paragraph_format.space_after=Pt(16)
doc.styles['Heading 2'].font.size=Pt(14);doc.styles['Heading 2'].paragraph_format.space_before=Pt(10)
doc.styles['Caption'].font.size=Pt(8.5);doc.styles['Caption'].font.italic=False;doc.styles['Caption'].font.bold=False
sec.different_first_page_header_footer=True
h=sec.header.paragraphs[0];h.text='CARENOW     EXHIBITION REPORT';h.runs[0].font.size=Pt(8)
f=sec.footer.paragraphs[0];f.text='CareNow 1.2.1  /  5 September 2026';f.runs[0].font.size=Pt(8);f.paragraph_format.tab_stops.clear_all();f.paragraph_format.tab_stops.add_tab_stop(Inches(6.79),WD_TAB_ALIGNMENT.RIGHT);f.add_run('\t')
fld=OxmlElement('w:fldSimple');fld.set(qn('w:instr'),'PAGE');f._p.append(fld)
doc.core_properties.title='CareNow Exhibition Report';doc.core_properties.subject='Product evidence market research and pilot proposal';doc.core_properties.author='CareNow project';doc.core_properties.keywords='CareNow, Bangladesh, Innovation Exhibitor, Expo, Hono, D1'
def p(s,style=None):return doc.add_paragraph(s,style)
def heading(s):doc.add_heading(s,level=1)
def page():doc.add_page_break()
def table(rows):
 t=doc.add_table(rows=1,cols=len(rows[0]));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
 cols=len(rows[0]); widths=([4.65,2.14] if cols==2 else [1.6,2.5,2.69] if cols==3 else [1.85,1.35,1.8,1.79]);
 if cols==3 and rows[0][0] in ['Layer','Organisation','Risk','Stage','Capability','Measure','Audience']: widths=[1.38,2.57,2.84]
 for i,w in enumerate(widths):t.columns[i].width=Inches(w)
 for ri,row in enumerate(rows):
  cells=t.rows[0].cells if ri==0 else t.add_row().cells
  for ci,s in enumerate(row):
   c=cells[ci];c.width=Inches(widths[ci]);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
   pp=c.paragraphs[0];pp.paragraph_format.space_after=Pt(0);pp.paragraph_format.line_spacing=1.05
   
   if ri>0 and re.match(r'^(?:BDT |USD )?\d',str(s)) and ci>0:pp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
   rr=pp.add_run(str(s));rr.font.size=Pt(9);rr.bold=ri==0;rr.font.color.rgb=RGBColor.from_string('FFFFFF' if ri==0 else '000000')
   tcPr=c._tc.get_or_add_tcPr();shade=OxmlElement('w:shd');shade.set(qn('w:fill'),'333333' if ri==0 else ('F4F5F5' if ri%2 else 'FFFFFF'));tcPr.append(shade)
   mar=OxmlElement('w:tcMar')
   for side in ['top','left','bottom','right']:
    z=OxmlElement('w:'+side);z.set(qn('w:w'),'95');z.set(qn('w:type'),'dxa');mar.append(z)
   tcPr.append(mar);borders=OxmlElement('w:tcBorders')
   for side in ['top','left','bottom','right']:
    z=OxmlElement('w:'+side);z.set(qn('w:val'),'single');z.set(qn('w:sz'),'4');z.set(qn('w:color'),'D9D9D9');borders.append(z)
   tcPr.append(borders)
  pr=t.rows[ri]._tr.get_or_add_trPr();pr.append(OxmlElement('w:cantSplit'))
  if ri==0:pr.append(OxmlElement('w:tblHeader'))
 doc.add_paragraph().paragraph_format.space_after=Pt(0)
 return t
# Cover
p('CareNow Exhibition Report','Title');p('Family care and local transport in Bangladesh','Subtitle');p('Product evidence  •  Market analysis  •  Pilot proposal')
doc.add_picture(str(R/'docs/report-assets/family-care-editorial.png'),width=Inches(6.79))
p('Original AI-generated editorial illustration. No real customer or care provider is depicted.','Caption')
p('Prepared for judges, researchers, public leaders, investors and technical reviewers. This report presents a working coordination prototype and the evidence needed to assess a responsible pilot. Its main recommendation is a bounded evaluation before real service expansion.')
p('Innovator Exhibitor  /  Bangladesh Innovation Fair context\nCareNow project  /  Edition 1.2.1  /  5 September 2026')
p('Live demonstration  https://carenow-demo.pages.dev\nSource and Android release  https://github.com/Seyamalam/CareNow','Caption')
# Reading map with verified one-page chapter pagination.
page();heading('Contents')
p('A working product, an evidence review and a bounded pilot proposal. The product atlas contains 33 application screenshots.')
for idx,item in enumerate(data['pages']):
 pp=p(item['title']);pp.paragraph_format.space_after=Pt(5);pp.paragraph_format.tab_stops.add_tab_stop(Inches(6.79),WD_TAB_ALIGNMENT.RIGHT);pp.add_run('\t'+str(idx+3))
for label,num in [('Product atlas',28),('References',39),('Glossary',42)]:
 pp=p(label);pp.paragraph_format.space_after=Pt(5);pp.paragraph_format.tab_stops.add_tab_stop(Inches(6.79),WD_TAB_ALIGNMENT.RIGHT);pp.add_run('\t'+str(num))
# sections, each starts deliberately. No forced row heights.
for index,item in enumerate(data['pages']):
 page();heading(item['title'])
 for s in item['paras']:p(s)
 if item['table']:table(item['table'])
 if item['title']=='Product positioning and business model':
  doc.add_picture(str(R/'docs/report-assets/neighbourhood-transport-editorial.png'),width=Inches(6.79));p('AI-generated editorial scene illustrating the broader transport concept. It is not evidence of a fleet or service operation.','Caption')
# Atlas with three full screenshots aligned in one paragraph; no layout tables.
for i,(title,imgs,body) in enumerate(data['atlas'],1):
 page();heading(title);p(f'Product atlas {i:02}  /  Fictional demonstration data','Caption')
 pp=doc.add_paragraph();pp.paragraph_format.space_after=Pt(12);pp.paragraph_format.line_spacing=1;pp.paragraph_format.keep_together=True
 for j,(name,cap) in enumerate(imgs):
  fn=R/'docs/screenshots'/name
  if not fn.exists():fn=R/'docs/report-assets'/name
  # Equal width preserves original images and full viewport. Spaces form gutters.
  pp.add_run().add_picture(str(fn),width=Inches(1.98))
  if j<2:pp.add_run(' ')
 p('Left: '+imgs[0][1]+'.  Centre: '+imgs[1][1]+'.  Right: '+imgs[2][1]+'.','Caption')
 p(body)
 p('Source: CareNow application capture, 5 September 2026. Native images use the corrected release. Web images use the same application. Screens with a scrollable body show the visible viewport.','Caption')
# Source register, 5-6 references per page; long URLs remain linked via Word auto-recognition.
for start,end in [(0,6),(6,12),(12,16)]:
 page();heading('References' if start==0 else 'References continued')
 p('Public sources accessed on 5 September 2026. Bracketed numbers in the report identify these entries. Publication and observation dates are retained where relevant.','Caption')
 for i in range(start,end):
  org,title,note,url=data['sources'][i]
  pp=p(f'[{i+1}] {org}  {title}');pp.runs[0].bold=True
  p(note)
  pp=p(url,'Caption');pp.paragraph_format.space_after=Pt(14)
 if start==12:
  doc.add_heading('Project and image provenance',level=2)
  p('The private Project DBA source document is summarised in the repository visual brief review and is not reproduced here. Native screenshots 61 to 82 show the layout release. Web captures show the deployed demo. Original illustrations were generated using the built-in image tool and are stored with prompts in the project assets.')
  p('Commercial assumptions, budget allocations and proposed pilot thresholds are project planning inputs. They are not externally sourced forecasts. No claim of clinical efficacy, government endorsement or existing commercial traction is made.')
page();heading('Glossary')
for term,definition in [
 ('API','Application programming interface. The contract through which the app requests data and actions from the server.'),
 ('BDT','Bangladeshi taka. Lakh denotes 100,000; one million taka equals ten lakh.'),
 ('CAC','Customer acquisition cost. The illustrative model allocates BDT 600 to acquiring one household.'),
 ('Contribution','Platform fee less the variable costs associated with completing a booking. It excludes fixed costs and tax.'),
 ('D1','Cloudflare’s managed SQL database. The demo stores versioned, isolated session data and attachments.'),
 ('GMV','Gross merchandise value. Total booking value before provider payments. It is not platform revenue.'),
 ('GPS','Global positioning system. The prototype animates saved routes; it does not dispatch or track a real vehicle.'),
 ('Hono','The web framework used for the Cloudflare Worker API.'),
 ('MapLibre','The map rendering library. Map rendering, geographic data and road routing are separate components.'),
 ('OOP','Out-of-pocket household health expenditure. The cited share is a historical national indicator.'),
 ('Pilot','A bounded evaluation under an agreed protocol. The proposed pilot has not been conducted.'),
 ('RBAC','Role-based access control. Production access also needs checks for the particular recipient, job and organisation.'),
 ('Take rate','The platform fee expressed as a share of booking value. Twelve percent is a modelling assumption.'),
 ('Type safety','Compile-time contracts combined here with runtime Zod validation. It does not by itself establish security or service safety.')]:
 pp=p(term);pp.paragraph_format.space_after=Pt(2);pp.runs[0].bold=True
 pp=p(definition);pp.paragraph_format.space_after=Pt(8)
# Remove inherited decorative paragraph borders and theme-font overrides.
for element in [doc.styles.element,doc.element]:
 for border in element.xpath('.//w:pBdr'):border.getparent().remove(border)
 for fonts in element.xpath('.//w:rFonts'):
  for key in list(fonts.attrib):
   if 'theme' in key.lower():del fonts.attrib[key]
# Save
out=R/'deliverables/CareNow-Exhibition-Report.docx';doc.save(out)
print(out)
